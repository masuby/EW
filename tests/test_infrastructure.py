"""Tests for builder infrastructure — generates test documents for visual verification."""

import sys
sys.path.insert(0, "/home/kaijage/model/EW")

from pathlib import Path
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

from src.builders.styles import (
    COLOR_RED_MAJOR, COLOR_ORANGE_WARNING, COLOR_YELLOW_ADVISORY,
    HEX_RED_MAJOR, HEX_ORANGE_WARNING, HEX_YELLOW_ADVISORY,
    HEX_DARK_TEAL, HEX_DARK_BLUE, HEX_LIGHT_BLUE_BG,
    FONT_PRIMARY, SIZE_BODY, SIZE_LABEL,
)
from src.builders.table_helpers import (
    set_cell_shading, set_cell_borders, remove_all_borders,
    set_paragraph_shading, create_no_border_table,
)
from src.builders.header_helpers import build_722e4_header, build_multirisk_header

OUTPUT_DIR = Path("/home/kaijage/model/EW/output")
OUTPUT_DIR.mkdir(exist_ok=True)


def test_table_helpers():
    """Test cell shading, borders, and paragraph shading."""
    doc = Document()

    # Test 1: Cell shading with alert colors
    p = doc.add_paragraph("Test 1: Alert Level Cell Colors")
    p.runs[0].bold = True

    table = doc.add_table(rows=1, cols=3)
    colors = [
        (HEX_RED_MAJOR, "MAJOR WARNING", "FFFFFF"),
        (HEX_ORANGE_WARNING, "WARNING", "FFFFFF"),
        (HEX_YELLOW_ADVISORY, "ADVISORY", "000000"),
    ]
    for i, (bg_hex, label, text_hex) in enumerate(colors):
        cell = table.cell(0, i)
        set_cell_shading(cell, bg_hex)
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(label)
        run.font.name = FONT_PRIMARY
        run.font.size = SIZE_LABEL
        run.bold = True
        run.font.color.rgb = RGBColor.from_string(text_hex)

    # Test 2: Hazard panel header colors
    doc.add_paragraph()
    p = doc.add_paragraph("Test 2: Hazard Panel Header Colors")
    p.runs[0].bold = True

    table2 = doc.add_table(rows=1, cols=4)
    hazards = [
        (HEX_DARK_TEAL, "Heavy Rain"),
        (HEX_DARK_TEAL, "Large Waves"),
        (HEX_DARK_TEAL, "Strong Winds"),
        (HEX_DARK_BLUE, "Floods"),
    ]
    for i, (bg_hex, label) in enumerate(hazards):
        cell = table2.cell(0, i)
        set_cell_shading(cell, bg_hex)
        p = cell.paragraphs[0]
        run = p.add_run(label)
        run.font.name = FONT_PRIMARY
        run.font.size = SIZE_LABEL
        run.bold = True
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)

    # Test 3: Paragraph shading (section heading bar)
    doc.add_paragraph()
    p = doc.add_paragraph("Test 3: Section Heading Bar")
    p.runs[0].bold = True

    p_bar = doc.add_paragraph()
    set_paragraph_shading(p_bar, HEX_LIGHT_BLUE_BG)
    run = p_bar.add_run("Multi-hazard assessment and recommendations")
    run.font.name = FONT_PRIMARY
    run.font.size = SIZE_BODY
    run.bold = True

    # Test 4: Borderless table
    doc.add_paragraph()
    p = doc.add_paragraph("Test 4: Borderless Table")
    p.runs[0].bold = True

    table3 = create_no_border_table(doc, 2, 2)
    table3.cell(0, 0).text = "Top-Left"
    table3.cell(0, 1).text = "Top-Right"
    table3.cell(1, 0).text = "Bottom-Left"
    table3.cell(1, 1).text = "Bottom-Right"

    out_path = OUTPUT_DIR / "test_table_helpers.docx"
    doc.save(str(out_path))
    assert out_path.exists()
    print(f"  Saved: {out_path}")


def test_722e4_header():
    """Test 722E_4 header generation."""
    doc = Document()

    # Set up page
    section = doc.sections[0]
    section.page_width = Cm(21.0)
    section.page_height = Cm(29.7)
    section.left_margin = Cm(1.27)
    section.right_margin = Cm(1.27)

    build_722e4_header(doc, include_issue_line=True,
                       issue_text="Issued on Saturday: 08-03-2025: 15:30 (EAT)")

    out_path = OUTPUT_DIR / "test_722e4_header.docx"
    doc.save(str(out_path))
    assert out_path.exists()
    print(f"  Saved: {out_path}")


def test_multirisk_header_en_new():
    """Test Multirisk header — English new format."""
    doc = Document()
    section = doc.sections[0]
    section.page_width = Cm(21.0)
    section.page_height = Cm(29.7)
    section.left_margin = Cm(1.27)
    section.right_margin = Cm(1.27)

    build_multirisk_header(doc, lang="en", header_variant="new",
                           bulletin_number=79,
                           issue_date_str="12/03/2025")

    out_path = OUTPUT_DIR / "test_multirisk_header_en_new.docx"
    doc.save(str(out_path))
    assert out_path.exists()
    print(f"  Saved: {out_path}")


def test_multirisk_header_sw():
    """Test Multirisk header — Swahili format."""
    doc = Document()
    section = doc.sections[0]
    section.page_width = Cm(21.0)
    section.page_height = Cm(29.7)
    section.left_margin = Cm(1.27)
    section.right_margin = Cm(1.27)

    build_multirisk_header(doc, lang="sw", header_variant="new",
                           bulletin_number=122,
                           issue_date_str="07/05/2025",
                           issue_time_str="9:42")

    out_path = OUTPUT_DIR / "test_multirisk_header_sw.docx"
    doc.save(str(out_path))
    assert out_path.exists()
    print(f"  Saved: {out_path}")


if __name__ == "__main__":
    print("Running infrastructure tests...")
    test_table_helpers()
    test_722e4_header()
    test_multirisk_header_en_new()
    test_multirisk_header_sw()
    print("All infrastructure tests passed!")
