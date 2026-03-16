"""Builder for Tanzania Multirisk Three Days Impact-Based Forecast Bulletin."""

from datetime import date
from pathlib import Path

from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT

from .base_builder import BaseBulletinBuilder
from .styles import (
    COLOR_BLUE_TMA, COLOR_BLACK, COLOR_WHITE, COLOR_DARK_GRAY,
    HEX_RED_MAJOR, HEX_ORANGE_WARNING, HEX_YELLOW_ADVISORY,
    HEX_DARK_TEAL, HEX_DARK_BLUE, HEX_LIGHT_BLUE_BG, HEX_WHITE,
    ALERT_COLORS, ALERT_TEXT_COLORS, HAZARD_HEADER_COLORS,
    FONT_PRIMARY,
    SIZE_TITLE_LARGE, SIZE_TITLE, SIZE_SUBTITLE, SIZE_BODY,
    SIZE_BODY_SMALL, SIZE_LABEL, SIZE_SMALL, SIZE_FOOTER,
    MAP_PANEL_WIDTH, MAP_PANEL_HEIGHT,
    MAP_SUMMARY_WIDTH, MAP_SUMMARY_HEIGHT,
    LOGO_PARTNER_WIDTH, LOGO_PARTNER_HEIGHT,
    MARGIN_LEFT, MARGIN_RIGHT, PAGE_WIDTH,
)
from .table_helpers import (
    set_cell_shading, set_cell_borders, remove_all_borders,
    set_cell_vertical_alignment, set_cell_margins, set_cell_width,
    set_paragraph_shading, create_no_border_table,
    set_table_fixed_layout, set_table_width, set_table_col_widths,
)
from .header_helpers import build_multirisk_header, _add_run
from ..models.common import AlertLevel, HazardType, Language
from ..models.multirisk import (
    MultiriskBulletin, MultiriskDayForecast, DaySummary,
    AlertTierEntry, TmaComment, MowComment, DmdComment,
)
from ..i18n.translations import t

ASSETS_DIR = Path(__file__).parent.parent.parent / "assets"

# Hazard type ordering for the 4-panel display
HAZARD_ORDER = [
    HazardType.HEAVY_RAIN,
    HazardType.LARGE_WAVES,
    HazardType.STRONG_WIND,
    HazardType.FLOODS,
]

HAZARD_LABEL_KEYS = {
    HazardType.HEAVY_RAIN: "hazard_heavy_rain",
    HazardType.LARGE_WAVES: "hazard_large_waves",
    HazardType.STRONG_WIND: "hazard_strong_winds",
    HazardType.FLOODS: "hazard_floods",
    HazardType.LANDSLIDES: "hazard_landslides",
    HazardType.EXTREME_TEMPERATURE: "hazard_extreme_temperature",
}

HAZARD_BG_COLORS = {
    HazardType.HEAVY_RAIN: HEX_DARK_TEAL,
    HazardType.LARGE_WAVES: HEX_DARK_TEAL,
    HazardType.STRONG_WIND: HEX_DARK_TEAL,
    HazardType.FLOODS: HEX_DARK_BLUE,
    HazardType.LANDSLIDES: HEX_DARK_TEAL,
    HazardType.EXTREME_TEMPERATURE: HEX_DARK_TEAL,
}

HAZARD_ICON_FILES = {
    HazardType.HEAVY_RAIN: "heavy_rain_white.png",
    HazardType.LARGE_WAVES: "large_waves_white.png",
    HazardType.STRONG_WIND: "strong_wind_white.png",
    HazardType.FLOODS: "floods_white.png",
    HazardType.LANDSLIDES: "landslides_white.png",
    HazardType.EXTREME_TEMPERATURE: "extreme_temperature_white.png",
}


class MultiriskBuilder(BaseBulletinBuilder):
    """Builds the Tanzania Multirisk Three Days Impact-Based Forecast Bulletin."""

    def __init__(self, data: MultiriskBulletin, maps_dir: str = None):
        super().__init__(data, language=data.language.value, maps_dir=maps_dir)
        self.bulletin: MultiriskBulletin = data
        self.lang = data.language.value

    def _format_date(self, d: date) -> str:
        return d.strftime("%d-%m-%Y")

    def _format_date_slash(self, d: date) -> str:
        return d.strftime("%d/%m/%Y")

    def _add_header(self):
        """Add the Multirisk header."""
        build_multirisk_header(
            self.doc,
            lang=self.lang,
            header_variant=self.bulletin.header_variant,
            bulletin_number=self.bulletin.bulletin_number,
            issue_date_str=self._format_date_slash(self.bulletin.issue_date),
            issue_time_str=self.bulletin.issue_time.strftime("%H:%M") if self.lang == "en"
                else self.bulletin.issue_time.strftime("%-H:%M"),
        )

    def _build_day_heading(self, day: MultiriskDayForecast):
        """Add 'SIKU YA X, tarehe DD-MM-YYYY' or 'DAY X - DD-MM-YYYY' heading."""
        p = self.doc.add_paragraph()
        p.paragraph_format.space_before = Pt(8)
        p.paragraph_format.space_after = Pt(6)
        heading_text = t("day_heading", self.lang,
                         n=day.day_number,
                         date=self._format_date(day.forecast_date))
        _add_run(p, heading_text, size=SIZE_TITLE_LARGE, bold=True,
                 color=COLOR_DARK_GRAY)

    def _build_hazard_panels(self, day: MultiriskDayForecast):
        """Build the 4 hazard map panels in a row."""
        # 2-row table: row 0 = colored headers, row 1 = maps
        table = self.doc.add_table(rows=2, cols=4)
        table.autofit = False
        table.alignment = WD_TABLE_ALIGNMENT.CENTER

        # Fixed layout: each column ~4.6cm = 2617 twips, total ~18.5cm
        col_w = 2617
        set_table_fixed_layout(table)
        set_table_col_widths(table, [col_w, col_w, col_w, col_w])
        set_table_width(table, col_w * 4)

        # Build a lookup from hazard type to map image
        panel_maps = {}
        for panel in day.hazard_panels:
            panel_maps[panel.hazard_type] = panel.map_image

        for col_idx, hazard_type in enumerate(HAZARD_ORDER):
            # Row 0: Colored header
            header_cell = table.cell(0, col_idx)
            bg_hex = HAZARD_BG_COLORS[hazard_type]
            set_cell_shading(header_cell, bg_hex)
            set_cell_margins(header_cell, top=40, bottom=40, left=40, right=40)

            p = header_cell.paragraphs[0]
            p.paragraph_format.space_before = Pt(2)
            p.paragraph_format.space_after = Pt(2)
            label_key = HAZARD_LABEL_KEYS[hazard_type]
            label = t(label_key, self.lang)
            _add_run(p, label, size=SIZE_LABEL, bold=True, color=COLOR_WHITE)

            # Add hazard icon (right side of header)
            icon_file = HAZARD_ICON_FILES.get(hazard_type)
            if icon_file:
                icon_path = ASSETS_DIR / "icons" / icon_file
                if icon_path.exists():
                    _add_run(p, "  ", size=SIZE_LABEL)  # spacer
                    run = p.add_run()
                    run.add_picture(str(icon_path), width=Cm(0.45), height=Cm(0.45))

            # Row 1: Map image
            map_cell = table.cell(1, col_idx)
            p = map_cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER

            map_img = panel_maps.get(hazard_type)
            self._add_image_or_placeholder(p, map_img, MAP_PANEL_WIDTH, MAP_PANEL_HEIGHT)

            # Borders
            thin = {'val': 'single', 'sz': 4, 'color': 'CCCCCC'}
            set_cell_borders(header_cell, top=thin, bottom=thin, left=thin, right=thin)
            set_cell_borders(map_cell, top=thin, bottom=thin, left=thin, right=thin)

    def _build_assessment_section(self, day: MultiriskDayForecast):
        """Build the multi-hazard assessment section with tier labels and recommendations."""
        # Section heading bar
        p = self.doc.add_paragraph()
        p.paragraph_format.space_before = Pt(8)
        p.paragraph_format.space_after = Pt(4)
        set_paragraph_shading(p, HEX_LIGHT_BLUE_BG)
        heading = t("multi_hazard_title", self.lang)
        _add_run(p, heading, size=SIZE_BODY, bold=True)

        # 2-column layout: [Summary Map] | [Alert Tiers + Recommendations]
        table = create_no_border_table(self.doc, 1, 2)
        table.autofit = False

        # Fixed layout: left ~8cm for map, right ~10.5cm for tiers
        col_left = 4536   # ~8.0cm
        col_right = 5932  # ~10.5cm
        set_table_fixed_layout(table)
        set_table_col_widths(table, [col_left, col_right])
        set_table_width(table, col_left + col_right)

        # Left: Summary map
        left_cell = table.cell(0, 0)
        set_cell_vertical_alignment(left_cell, "top")
        p = left_cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        self._add_image_or_placeholder(p, day.summary_map,
                                        MAP_SUMMARY_WIDTH, MAP_SUMMARY_HEIGHT)

        # Right: Alert tier labels + recommendations
        right_cell = table.cell(0, 1)
        set_cell_vertical_alignment(right_cell, "top")
        right_cell.paragraphs[0].clear()

        # Three tier labels
        tier_data = {tier.alert_level: tier for tier in day.alert_tiers}
        tiers_to_show = [
            (AlertLevel.MAJOR_WARNING, "major_warning", HEX_RED_MAJOR, COLOR_WHITE),
            (AlertLevel.WARNING, "warning", HEX_ORANGE_WARNING, COLOR_WHITE),
            (AlertLevel.ADVISORY, "advisory", HEX_YELLOW_ADVISORY, COLOR_BLACK),
        ]

        first_para = True
        for alert_level, label_key, bg_hex, text_color in tiers_to_show:
            tier = tier_data.get(alert_level)
            none_text = t("tier_none", self.lang)

            # Create a small inner table for the label
            inner_table = self.doc.add_table(rows=1, cols=2)
            # Move inner table into right cell
            right_cell._tc.append(inner_table._tbl)

            inner_table.autofit = False

            # Label cell
            label_cell = inner_table.cell(0, 0)
            set_cell_shading(label_cell, bg_hex)
            set_cell_margins(label_cell, top=20, bottom=20, left=40, right=40)
            p = label_cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            label_text = t(label_key, self.lang)
            _add_run(p, label_text, size=SIZE_LABEL, bold=True, color=text_color)

            # Content cell
            content_cell = inner_table.cell(0, 1)
            remove_all_borders(content_cell)
            p = content_cell.paragraphs[0]

            if tier and tier.text:
                _add_run(p, f"  {tier.text}", size=SIZE_BODY_SMALL)
            elif tier is None or not tier.recommendations:
                _add_run(p, f"  {none_text}", size=SIZE_BODY_SMALL)

            # Remove borders from label cell too (visual cleanliness)
            remove_all_borders(label_cell)

        # Recommendations (if ADVISORY has them)
        advisory_tier = tier_data.get(AlertLevel.ADVISORY)
        recs = day.recommendations or (advisory_tier.recommendations if advisory_tier else [])

        if recs:
            # Recommendation intro
            intro = day.recommendation_intro
            if intro:
                p_intro = right_cell.add_paragraph()
                p_intro.paragraph_format.space_before = Pt(4)
                p_intro.paragraph_format.space_after = Pt(2)
                _add_run(p_intro, intro, size=SIZE_BODY_SMALL, bold=True)

            # Bullet points
            for rec in recs:
                p_rec = right_cell.add_paragraph()
                p_rec.paragraph_format.space_before = Pt(1)
                p_rec.paragraph_format.space_after = Pt(1)
                p_rec.paragraph_format.left_indent = Cm(0.5)
                # Bullet character
                _add_run(p_rec, "\u2022 ", size=SIZE_BODY_SMALL, bold=True)
                _add_run(p_rec, rec, size=SIZE_BODY_SMALL)

        # Committee note (bold paragraph)
        committee = day.committee_note
        if committee:
            p_comm = right_cell.add_paragraph()
            p_comm.paragraph_format.space_before = Pt(6)
            p_comm.paragraph_format.space_after = Pt(4)
            # Split into bold prefix and normal text
            _add_run(p_comm, committee, size=SIZE_BODY_SMALL, bold=True)

    def _build_outlook_page(self, day: MultiriskDayForecast):
        """Build the outlook/comments page for a single day."""
        # Outlook heading bar
        p = self.doc.add_paragraph()
        p.paragraph_format.space_before = Pt(8)
        p.paragraph_format.space_after = Pt(6)
        set_paragraph_shading(p, HEX_LIGHT_BLUE_BG)
        heading = t("outlook_heading", self.lang,
                     n=day.day_number,
                     date=self._format_date(day.forecast_date))
        _add_run(p, heading, size=SIZE_BODY, bold=True)

        # TMA Comments
        self._build_tma_section(day)

        # MoW Comments
        self._build_mow_section(day)

        # DMD Comments
        self._build_dmd_section(day)

    def _build_tma_section(self, day: MultiriskDayForecast):
        """Build the TMA expert analysis section."""
        p = self.doc.add_paragraph()
        p.paragraph_format.space_before = Pt(6)
        p.paragraph_format.space_after = Pt(4)
        _add_run(p, t("comments_tma", self.lang), size=SIZE_SUBTITLE, bold=True)

        if not day.tma_comment or not day.tma_comment.entries:
            p = self.doc.add_paragraph()
            _add_run(p, t("no_warning_text", self.lang), size=SIZE_BODY)
            return

        for entry in day.tma_comment.entries:
            # Description
            p = self.doc.add_paragraph()
            p.paragraph_format.space_before = Pt(2)
            p.paragraph_format.space_after = Pt(2)
            _add_run(p, entry.description, size=SIZE_BODY)

            # Likelihood/Impact
            if entry.rating:
                p = self.doc.add_paragraph()
                p.paragraph_format.space_before = Pt(4)
                p.paragraph_format.space_after = Pt(0)
                likelihood_val = t(f"level_{entry.rating.likelihood.value.lower()}", self.lang)
                _add_run(p, f"{t('likelihood', self.lang)}: {likelihood_val}",
                         size=SIZE_BODY)

                p = self.doc.add_paragraph()
                p.paragraph_format.space_before = Pt(0)
                p.paragraph_format.space_after = Pt(4)
                impact_val = t(f"level_{entry.rating.impact.value.lower()}", self.lang)
                _add_run(p, f"{t('impact', self.lang)}: {impact_val}",
                         size=SIZE_BODY)

    def _build_mow_section(self, day: MultiriskDayForecast):
        """Build the Ministry of Water section."""
        p = self.doc.add_paragraph()
        p.paragraph_format.space_before = Pt(6)
        p.paragraph_format.space_after = Pt(4)
        _add_run(p, t("comments_mow", self.lang), size=SIZE_SUBTITLE, bold=True)

        if not day.mow_comment or not day.mow_comment.entries:
            return

        for entry in day.mow_comment.entries:
            p = self.doc.add_paragraph()
            p.paragraph_format.space_before = Pt(2)
            p.paragraph_format.space_after = Pt(2)
            _add_run(p, entry.description, size=SIZE_BODY)

            if entry.rating:
                p = self.doc.add_paragraph()
                p.paragraph_format.space_before = Pt(4)
                p.paragraph_format.space_after = Pt(0)
                likelihood_val = t(f"level_{entry.rating.likelihood.value.lower()}", self.lang)
                _add_run(p, f"{t('likelihood', self.lang)}: {likelihood_val}",
                         size=SIZE_BODY)

                p = self.doc.add_paragraph()
                p.paragraph_format.space_before = Pt(0)
                p.paragraph_format.space_after = Pt(4)
                impact_val = t(f"level_{entry.rating.impact.value.lower()}", self.lang)
                _add_run(p, f"{t('impact', self.lang)}: {impact_val}",
                         size=SIZE_BODY)

    def _build_dmd_section(self, day: MultiriskDayForecast):
        """Build the DMD impact assessment section."""
        p = self.doc.add_paragraph()
        p.paragraph_format.space_before = Pt(6)
        p.paragraph_format.space_after = Pt(4)
        _add_run(p, t("comments_dmd", self.lang), size=SIZE_SUBTITLE, bold=True)

        if not day.dmd_comment:
            return

        # Header text
        if day.dmd_comment.header_text:
            p = self.doc.add_paragraph()
            p.paragraph_format.space_before = Pt(2)
            p.paragraph_format.space_after = Pt(2)
            _add_run(p, day.dmd_comment.header_text, size=SIZE_BODY, bold=True)

        # Impact bullets
        for bullet in day.dmd_comment.impact_bullets:
            p = self.doc.add_paragraph()
            p.paragraph_format.space_before = Pt(1)
            p.paragraph_format.space_after = Pt(1)
            p.paragraph_format.left_indent = Cm(0.5)
            _add_run(p, f"- {bullet}", size=SIZE_BODY)

        # Rating
        if day.dmd_comment.rating:
            p = self.doc.add_paragraph()
            p.paragraph_format.space_before = Pt(4)
            p.paragraph_format.space_after = Pt(0)
            likelihood_val = t(f"level_{day.dmd_comment.rating.likelihood.value.lower()}", self.lang)
            _add_run(p, f"{t('likelihood', self.lang)}: {likelihood_val}",
                     size=SIZE_BODY)

            p = self.doc.add_paragraph()
            p.paragraph_format.space_before = Pt(0)
            p.paragraph_format.space_after = Pt(4)
            impact_val = t(f"level_{day.dmd_comment.rating.impact.value.lower()}", self.lang)
            _add_run(p, f"{t('impact', self.lang)}: {impact_val}",
                     size=SIZE_BODY)

    def _build_summary_page(self):
        """Build the district summary page."""
        for summary in self.bulletin.day_summaries:
            # Day heading bar
            p = self.doc.add_paragraph()
            p.paragraph_format.space_before = Pt(10)
            p.paragraph_format.space_after = Pt(6)
            set_paragraph_shading(p, HEX_LIGHT_BLUE_BG)
            heading = t("summary_heading", self.lang, n=summary.day_number)
            _add_run(p, heading, size=SIZE_BODY, bold=True)

            # Three tier labels with district lists
            tiers = [
                ("major_warning", HEX_RED_MAJOR, COLOR_WHITE, summary.major_warning_districts),
                ("warning", HEX_ORANGE_WARNING, COLOR_WHITE, summary.warning_districts),
                ("advisory", HEX_YELLOW_ADVISORY, COLOR_BLACK, summary.advisory_districts),
            ]

            for label_key, bg_hex, text_color, districts in tiers:
                # Create tier row as a table
                tier_table = self.doc.add_table(rows=1, cols=2)
                tier_table.autofit = False

                # Label cell
                label_cell = tier_table.cell(0, 0)
                set_cell_shading(label_cell, bg_hex)
                set_cell_margins(label_cell, top=20, bottom=20, left=40, right=40)
                remove_all_borders(label_cell)
                p = label_cell.paragraphs[0]
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                label_text = t(label_key, self.lang)
                _add_run(p, label_text, size=SIZE_LABEL, bold=True, color=text_color)

                # District list cell
                dist_cell = tier_table.cell(0, 1)
                remove_all_borders(dist_cell)
                p = dist_cell.paragraphs[0]
                p.paragraph_format.left_indent = Cm(0.3)

                if districts:
                    district_text = ", ".join(districts)
                    _add_run(p, district_text, size=SIZE_BODY_SMALL)

    def _build_footer_page(self):
        """Build the footer with contact info and partner logos."""
        # Contact line
        p = self.doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_before = Pt(12)
        p.paragraph_format.space_after = Pt(6)
        _add_run(p, t("contact_line", self.lang), size=SIZE_BODY_SMALL)

        # Attribution
        p = self.doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_before = Pt(2)
        p.paragraph_format.space_after = Pt(12)
        _add_run(p, t("attribution", self.lang), size=SIZE_SMALL)

        # Partner logos
        logo_table = self.doc.add_table(rows=1, cols=4)
        logo_table.alignment = WD_TABLE_ALIGNMENT.CENTER

        logo_files = [
            "italian_ministry.png",
            "italian_agency.png",
            "undrr.png",
            "cima.png",
        ]

        for i, logo_file in enumerate(logo_files):
            cell = logo_table.cell(0, i)
            remove_all_borders(cell)
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER

            logo_path = ASSETS_DIR / "logos" / logo_file
            if logo_path.exists():
                run = p.add_run()
                run.add_picture(str(logo_path),
                               width=LOGO_PARTNER_WIDTH,
                               height=LOGO_PARTNER_HEIGHT)
            else:
                _add_run(p, f"[{logo_file}]", size=SIZE_SMALL)

    def build(self) -> Document:
        """Build the complete Multirisk bulletin document."""
        for day in self.bulletin.days:
            # Page A: Maps + Assessment
            self._add_header()
            self._build_day_heading(day)
            self._build_hazard_panels(day)
            self._build_assessment_section(day)

            # Page B: Outlook/Comments
            self._add_page_break()
            self._add_header()
            self._build_outlook_page(day)

            # Page break before next day (unless last day)
            if day.day_number < 3:
                self._add_page_break()

        # Summary page
        self._add_page_break()
        self._add_header()
        self._build_summary_page()

        # Footer
        self._build_footer_page()

        return self.doc
