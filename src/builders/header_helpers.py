"""Header construction utilities for both document types."""

from pathlib import Path

from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH

from .styles import (
    COLOR_BLUE_TMA, COLOR_BLACK, FONT_PRIMARY,
    SIZE_TITLE_LARGE, SIZE_TITLE, SIZE_SUBTITLE, SIZE_BODY,
    SIZE_BODY_SMALL, SIZE_SMALL, SIZE_LABEL,
    LOGO_COAT_WIDTH, LOGO_COAT_HEIGHT,
    LOGO_TMA_WIDTH, LOGO_TMA_HEIGHT,
)
from .table_helpers import (
    remove_all_borders, set_cell_vertical_alignment,
    set_table_fixed_layout, set_table_col_widths, set_table_width,
)
from ..i18n.translations import t, t_list

ASSETS_DIR = Path(__file__).parent.parent.parent / "assets"


def _add_run(paragraph, text, font_name=FONT_PRIMARY, size=SIZE_BODY,
             color=COLOR_BLACK, bold=False, italic=False):
    """Add a formatted run to a paragraph."""
    run = paragraph.add_run(text)
    run.font.name = font_name
    run.font.size = size
    run.font.color.rgb = color
    run.bold = bold
    run.italic = italic
    return run


def _try_add_image(cell, image_path, width, height):
    """Add image to cell if file exists, otherwise add placeholder text."""
    paragraph = cell.paragraphs[0]
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Try the given path, then try alternative extensions
    resolved = None
    if image_path:
        p = Path(image_path)
        if p.exists():
            resolved = p
        else:
            # Try alternative extensions
            for ext in ['.jpeg', '.jpg', '.png', '.bmp']:
                alt = p.with_suffix(ext)
                if alt.exists():
                    resolved = alt
                    break

    if resolved:
        run = paragraph.add_run()
        run.add_picture(str(resolved), width=width, height=height)
    else:
        _add_run(paragraph, "[Logo]", size=SIZE_SMALL, color=COLOR_BLACK)


def build_722e4_header(doc: Document, include_issue_line=True,
                       issue_text=None, total_width_twips: int = None):
    """Build the 722E_4 Five-Day Forecast header.

    Layout: [Coat of Arms] | [Title Text] | [TMA Logo]
    """
    table = doc.add_table(rows=1, cols=3)
    table.autofit = False

    # Fixed layout with explicit column widths (in twips: 1cm = 567 twips)
    coat_w = 1418   # 2.5cm
    tma_w = 1418    # 2.5cm
    if total_width_twips:
        title_w = total_width_twips - coat_w - tma_w
    else:
        title_w = 7654  # 13.5cm (portrait default)
    total_w = coat_w + title_w + tma_w
    set_table_fixed_layout(table)
    set_table_col_widths(table, [coat_w, title_w, tma_w])
    set_table_width(table, total_w)

    # Remove borders
    for cell in table.rows[0].cells:
        remove_all_borders(cell)
        set_cell_vertical_alignment(cell, "center")

    # Column 0: Coat of Arms
    coat_path = ASSETS_DIR / "logos" / "coat_of_arms.jpeg"
    _try_add_image(table.cell(0, 0), coat_path, LOGO_COAT_WIDTH, LOGO_COAT_HEIGHT)

    # Column 1: Title text (centered)
    center_cell = table.cell(0, 1)
    # Clear default paragraph
    center_cell.paragraphs[0].clear()

    # Line 1: Country name
    p = center_cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.line_spacing = Pt(14)
    _add_run(p, t("country_name", "en"), size=SIZE_TITLE, bold=True)

    # Line 2: Ministry
    p = center_cell.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.line_spacing = Pt(13)
    _add_run(p, t("ministry", "en"), size=SIZE_SUBTITLE, bold=True)

    # Line 3: TMA name (blue)
    p = center_cell.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.line_spacing = Pt(13)
    _add_run(p, t("tma_name", "en"), size=SIZE_SUBTITLE, bold=True,
             color=COLOR_BLUE_TMA)

    # Line 4: Five day title
    p = center_cell.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.line_spacing = Pt(14)
    _add_run(p, t("five_day_title", "en"), size=SIZE_TITLE, bold=True)

    # Column 2: TMA Logo
    tma_path = ASSETS_DIR / "logos" / "tma_logo.jpeg"
    _try_add_image(table.cell(0, 2), tma_path, LOGO_TMA_WIDTH, LOGO_TMA_HEIGHT)

    # Issue line (below header table)
    if include_issue_line and issue_text:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_before = Pt(2)
        p.paragraph_format.space_after = Pt(4)
        _add_run(p, issue_text, size=SIZE_TITLE, bold=True, color=COLOR_BLACK)

    return table


def build_multirisk_header(doc: Document, lang: str = "en",
                           header_variant: str = "new",
                           bulletin_number: int = 0,
                           issue_date_str: str = "",
                           issue_time_str: str = ""):
    """Build the Multirisk bulletin header.

    Layout: [Coat of Arms + Left Text] | [Right Institutional Text]
    """
    table = doc.add_table(rows=1, cols=3)
    table.autofit = False

    # Fixed layout: Coat=1.3cm, Left text=11.2cm, Right text=6cm = 18.5cm
    coat_w = 737    # 1.3cm
    left_w = 6350   # 11.2cm
    right_w = 3402  # 6cm
    set_table_fixed_layout(table)
    set_table_col_widths(table, [coat_w, left_w, right_w])
    set_table_width(table, coat_w + left_w + right_w)

    for cell in table.rows[0].cells:
        remove_all_borders(cell)
        set_cell_vertical_alignment(cell, "top")

    # Column 0: Coat of Arms (larger for multirisk)
    coat_path = ASSETS_DIR / "logos" / "coat_of_arms.jpeg"
    _try_add_image(table.cell(0, 0), coat_path, Cm(1.8), Cm(1.8))

    # Column 1: Left text block
    left_cell = table.cell(0, 1)
    left_cell.paragraphs[0].clear()

    # Country name
    p = left_cell.paragraphs[0]
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    if lang == "sw":
        _add_run(p, t("country_name", "sw"), size=SIZE_TITLE, bold=True)
    else:
        _add_run(p, t("mr_country", "en"), size=SIZE_TITLE, bold=True)

    # Subtitle
    p = left_cell.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    _add_run(p, t("mr_subtitle", lang), size=SIZE_BODY, bold=True,
             color=COLOR_BLUE_TMA)

    # Issue info
    p = left_cell.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    if lang == "sw":
        issue_text = t("mr_issued", "sw",
                       number=bulletin_number,
                       date=issue_date_str,
                       time=issue_time_str)
    else:
        issue_text = t("mr_issued", "en",
                       number=bulletin_number,
                       date=issue_date_str)
    _add_run(p, issue_text, size=SIZE_BODY_SMALL)

    # Column 2: Right institutional text
    right_cell = table.cell(0, 2)
    right_cell.paragraphs[0].clear()

    if lang == "sw":
        lines = t_list("mr_right_header", "sw")
    elif header_variant == "new":
        lines = t_list("mr_right_header_new", "en")
    else:
        lines = t_list("mr_right_header_old", "en")

    p = right_cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)

    for i, line in enumerate(lines):
        if i > 0:
            p = right_cell.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            p.paragraph_format.space_before = Pt(0)
            p.paragraph_format.space_after = Pt(0)
        is_bold = (i == 0)  # First line bold
        _add_run(p, line, size=SIZE_SMALL, bold=is_bold)

    return table
