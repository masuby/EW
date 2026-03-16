"""Abstract base document builder with shared functionality."""

from abc import ABC, abstractmethod
from pathlib import Path
from typing import Optional

from docx import Document
from docx.shared import Cm, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.section import WD_ORIENT

from .styles import (
    PAGE_WIDTH, PAGE_HEIGHT,
    MARGIN_TOP, MARGIN_BOTTOM, MARGIN_LEFT, MARGIN_RIGHT,
    COLOR_GRAY_PLACEHOLDER, FONT_PRIMARY, SIZE_BODY_SMALL,
    HEX_GRAY_PLACEHOLDER,
)
from .table_helpers import set_cell_shading, remove_all_borders


class BaseBulletinBuilder(ABC):
    """Abstract base class for bulletin document builders."""

    def __init__(self, data, language: str = "en", maps_dir: Optional[str] = None):
        self.data = data
        self.language = language
        self.maps_dir = Path(maps_dir) if maps_dir else None
        self.doc = Document()
        self._setup_page()

    def _setup_page(self):
        """Configure A4 page size and margins."""
        section = self.doc.sections[0]
        section.page_width = PAGE_WIDTH
        section.page_height = PAGE_HEIGHT
        section.top_margin = MARGIN_TOP
        section.bottom_margin = MARGIN_BOTTOM
        section.left_margin = MARGIN_LEFT
        section.right_margin = MARGIN_RIGHT
        section.orientation = WD_ORIENT.PORTRAIT

        # Subclasses can override to landscape
        self._apply_orientation()

        # Set default font
        style = self.doc.styles['Normal']
        font = style.font
        font.name = FONT_PRIMARY
        font.size = SIZE_BODY_SMALL

    def _resolve_image(self, image_ref) -> Optional[Path]:
        """Resolve a map image reference to an absolute file path."""
        if image_ref is None:
            return None

        # Handle MapImage dataclass
        file_path = getattr(image_ref, 'file_path', image_ref)
        if file_path is None:
            return None

        path = Path(file_path)
        if path.is_absolute() and path.exists():
            return path

        # Try relative to maps_dir
        if self.maps_dir:
            resolved = self.maps_dir / path
            if resolved.exists():
                return resolved

        # Try relative to project root
        project_root = Path(__file__).parent.parent.parent
        resolved = project_root / path
        if resolved.exists():
            return resolved

        return None

    def _add_image_or_placeholder(self, paragraph, image_ref, width, height):
        """Insert an image into a paragraph, or a placeholder if image is missing."""
        resolved = self._resolve_image(image_ref)
        if resolved:
            run = paragraph.add_run()
            run.add_picture(str(resolved), width=width, height=height)
        else:
            # Insert placeholder text
            run = paragraph.add_run("[Map]")
            run.font.name = FONT_PRIMARY
            run.font.size = SIZE_BODY_SMALL
            run.font.color.rgb = COLOR_GRAY_PLACEHOLDER

    def _add_map_placeholder_cell(self, cell, width, height):
        """Add a gray placeholder box in a table cell for missing maps."""
        set_cell_shading(cell, HEX_GRAY_PLACEHOLDER)
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run("\n\n[Map]\n\n")
        run.font.name = FONT_PRIMARY
        run.font.size = SIZE_BODY_SMALL

    def _apply_orientation(self):
        """Override in subclass to change orientation (called from _setup_page)."""
        pass

    def _set_landscape(self):
        """Switch first section to landscape A4."""
        section = self.doc.sections[0]
        section.orientation = WD_ORIENT.LANDSCAPE
        section.page_width = PAGE_HEIGHT   # swap: 29.7cm
        section.page_height = PAGE_WIDTH   # swap: 21.0cm

    def _add_page_break(self):
        """Add a page break to the document."""
        self.doc.add_page_break()

    def _add_section_break(self):
        """Add a new section (copying page setup from the first section)."""
        first = self.doc.sections[0]
        new_section = self.doc.add_section()
        new_section.page_width = first.page_width
        new_section.page_height = first.page_height
        new_section.orientation = first.orientation
        new_section.top_margin = MARGIN_TOP
        new_section.bottom_margin = MARGIN_BOTTOM
        new_section.left_margin = MARGIN_LEFT
        new_section.right_margin = MARGIN_RIGHT
        return new_section

    @abstractmethod
    def build(self) -> Document:
        """Build the complete document. Returns the Document object."""
        ...

    def save(self, output_path: str):
        """Save the document to disk."""
        output = Path(output_path)
        output.parent.mkdir(parents=True, exist_ok=True)
        self.doc.save(str(output))
        return str(output)
